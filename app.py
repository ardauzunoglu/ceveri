import os
import json
import time
import pandas as pd
import boto3
import six
import docx
from flask import Flask, flash, request, redirect, send_file, url_for, render_template, Response
from werkzeug.utils import secure_filename
from googletrans import Translator

ALLOWED_EXTENSIONS = ["TXT", "JSON", "CSV", "XLSX", "XML", "DOCX", "PKL"]
AWS_ACCESS_KEY_ID=os.environ.get('AWS_ACCESS_KEY_ID')
AWS_SECRET_KEY=os.environ.get('AWS_SECRET_KEY')

language_abbs = {
    "en":"İngilizce",
    "fr":"Fransızca",
    "es":"İspanyolca",
    "de":"Almanca",
    "ru":"Rusca",
    "ar":"Arapça",
    "it":"İtalyanca",
    "nl":"Flemenkçe",
    "ko":"Korece",
    "zh":"Çince"
}

app = Flask(__name__)
app.config['ALLOWED_EXTENSIONS'] = ALLOWED_EXTENSIONS
app.config['MAX_CONTENT_LENGTH'] = 1 * 1024 * 1024
s3_client = boto3.client('s3',
                        aws_access_key_id=AWS_ACCESS_KEY_ID,
                        aws_secret_access_key=AWS_SECRET_KEY)
translate_client = Translator()

def translate(extension, file, source_language, detect_source_language):
    chars_translated = 0
    source_languages = []
    def translate_text(target, text, source_language, detect_source_language):
        if isinstance(text, six.binary_type):
            text = text.decode("utf-8")

        if text == "":
            return ""

        if detect_source_language:
            source_language = translate_client.detect(text).lang
            result = translate_client.translate(text, src=source_language, dest=target)
            if source_language not in source_languages:
                source_languages.append(source_language)

        else:
            result = translate_client.translate(text, src=source_language, dest=target)

        return result.text
    
    def translate_json(json_file, source_language, chars_translated, detect_source_language):
        search_dict = json.load(open(json_file, "r", encoding="utf-8"))
        def translate_recursively(search_dict, field, chars_translated, detect_source_language):
            for key, value in search_dict.items():
                if isinstance(value, field):
                    if "https://" not in value:
                        translation = translate_text("tr", value, source_language, detect_source_language)
                        chars_translated += len(translation)
                        search_dict[key] = translation

                elif isinstance(value, dict):
                    results = translate_recursively(value, field, chars_translated, detect_source_language)

                elif isinstance(value, (list,tuple)):
                    for item in value:
                        if isinstance(item, dict):
                            more_results = translate_recursively(item, field, chars_translated, detect_source_language)

                        elif isinstance(item, field):
                            if "https://" not in item:
                                translation = translate_text("tr", item, source_language, detect_source_language)
                                chars_translated += len(translation)
                                value[value.index(item)] = translation

                    search_dict[key] = value
                        
            return search_dict, chars_translated
        
        translated_dict, chars_translated = translate_recursively(search_dict, str, chars_translated, detect_source_language)
        return translated_dict, chars_translated
    
    def translate_excel(excel_file, source_language, detect_source_language):
        df = pd.read_excel(excel_file)
        chars_translated = 0

        for index, row in list(df.iterrows()):
            for column in list(df.columns):
                if isinstance(row[column], str):
                    translation = translate_text("tr", row[column], source_language, detect_source_language)
                    df.replace(row[column], translation, inplace=True)
                    chars_translated += len(translation)

        return df, chars_translated

    def translate_csv(csv_file, source_language, detect_source_language):
        df = pd.read_csv(csv_file)
        chars_translated = 0

        for index, row in list(df.iterrows()):
            for column in list(df.columns):
                if isinstance(row[column], str):
                    translation = translate_text("tr", row[column], source_language, detect_source_language)
                    df.replace(row[column], translation, inplace=True)
                    chars_translated += len(translation)

        return df, chars_translated

    def translate_xml(xml_file, source_language, detect_source_language):
        df = pd.read_xml(xml_file)
        chars_translated = 0

        for index, row in list(df.iterrows()):
            for column in list(df.columns):
                if isinstance(row[column], str):
                    translation = translate_text("tr", row[column], source_language, detect_source_language)
                    df.replace(row[column], translation, inplace=True)
                    chars_translated += len(translation)

        return df, chars_translated

    def translate_pkl(pkl_file, source_language, detect_source_language):
        df = pd.read_pickle(pkl_file)
        chars_translated = 0

        for index, row in list(df.iterrows()):
            for column in list(df.columns):
                if isinstance(row[column], str):
                    translation = translate_text("tr", row[column], source_language, detect_source_language)
                    df.replace(row[column], translation, inplace=True)
                    chars_translated += len(translation)

        return df, chars_translated

    def translate_txt(txt_file, source_language, detect_source_language):
        data = open(txt_file, "r", encoding="utf-8").readlines()
        translations = ["\n"] * len(data)
        chars_translated = 0

        for datum in data:
            if datum != "\n":
                translation = translate_text("tr", datum.replace("\n", ""), source_language, detect_source_language)
                translations[data.index(datum)] = translation
                chars_translated += len(translation)
            else:
                translations[data.index(datum)] = datum

        return translations, chars_translated

    def translate_docx(docx_file, source_language, detect_source_language):
        doc = docx.Document(docx_file)
        translated_doc = docx.Document()
        chars_translated = 0

        paragraphs = doc.paragraphs
        for paragraph in paragraphs:
            chars_translated += len(paragraph.text)
            translated_paragraph = translate_text("tr", paragraph.text, source_language, detect_source_language)
            translated_doc.add_paragraph(translated_paragraph)

        return translated_doc, chars_translated

    if extension.lower() == "json":
        translated_dict, chars_translated = translate_json(file, source_language, chars_translated, detect_source_language)
        save_path = ".".join(file.split(".")[:-1]) + "_tr" + "." + extension 
        with open(save_path, 'w', encoding="utf-8") as f:
            json.dump(translated_dict, f, indent=4, ensure_ascii=False)
        s3_client.upload_file(save_path, "ceveri", save_path)

    if extension.lower() == "xlsx":
        translated_df, chars_translated = translate_excel(file, source_language, detect_source_language)
        save_path = ".".join(file.split(".")[:-1]) + "_tr" + "." + extension 
        translated_df.to_excel(save_path)
        s3_client.upload_file(save_path, "ceveri", save_path)

    if extension.lower() == "csv":
        translated_df, chars_translated = translate_csv(file, source_language, detect_source_language)
        save_path = ".".join(file.split(".")[:-1]) + "_tr" + "." + extension 
        translated_df.to_csv(save_path)
        s3_client.upload_file(save_path, "ceveri", save_path)

    if extension.lower() == "xml":
        translated_df, chars_translated = translate_xml(file, source_language, detect_source_language)
        save_path = ".".join(file.split(".")[:-1]) + "_tr" + "." + extension 
        translated_df.to_xml(save_path)
        s3_client.upload_file(save_path, "ceveri", save_path)

    if extension.lower() == "pkl":
        translated_df, chars_translated = translate_pkl(file, source_language, detect_source_language)
        save_path = ".".join(file.split(".")[:-1]) + "_tr" + "." + extension 
        translated_df.to_pkl(save_path)
        s3_client.upload_file(save_path, "ceveri", save_path)

    if extension.lower() == "txt":
        translation_list, chars_translated = translate_txt(file, source_language, detect_source_language)
        save_path = ".".join(file.split(".")[:-1]) + "_tr" + "." + extension 
        with open(save_path, "w", encoding="utf-8") as f:
            for line in translation_list:
                if line != "\n":
                    f.writelines(line+"\n")
                else:
                    f.writelines(line)
        s3_client.upload_file(save_path, "ceveri", save_path)

    if extension.lower() == "docx":
        translated_doc, chars_translated = translate_docx(file, source_language, detect_source_language)
        save_path = ".".join(file.split(".")[:-1]) + "_tr" + "." + extension 
        translated_doc.save(save_path)
        s3_client.upload_file(save_path, "ceveri", save_path)

    return save_path, chars_translated, source_languages

def is_allowed_ext(filename):
    if "." not in filename:
        return False
    
    extension = filename.rsplit(".", 1)[1]
    
    if extension.upper() in app.config['ALLOWED_EXTENSIONS']:
        return True
    else:
        return False

@app.errorhandler(413)
def size_too_big(e):
    error_headline = "Dosya boyutu çok büyük."
    error_text = "Yüklediğiniz dosyanın boyutu çok büyük. Anlık olarak yükleyebileceğiniz maksimum dosya boyutu 1MB'dir. Lütfen boyutu 1MB'nin altında olan başka bir dosya ile tekrar deneyiniz."
    return render_template("index.html", error_headline=error_headline, error_text=error_text)

@app.route("/anasayfa", methods=["GET", "POST"])
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        if request.files:
            file = request.files["data-file"]

            if file.filename == "":
                error_headline = "Dosya bir isme sahip olmalıdır."
                error_text = "Yüklediğiniz dosya bir isme sahip değil. Yüklediğiniz dosyanın işlenebilmesi için bir isme sahip olmalıdır, lütfen yeniden deneyiniz."
                return render_template("index.html", error_headline=error_headline, error_text=error_text)

            if not is_allowed_ext(file.filename):
                error_headline = "Dosya uzantısı desteklenmiyor."
                error_text = "Yüklediğiniz dosyanın sahip olduğu uzantı şu anlık desteklenmyior. Anlık olarak desteklenen uzantılar arasında .txt, .json, .xlsx, .csv uzantıları yer almakta. Lütfen desteklenen uzantılardan birine sahip başka bir dosya ile tekrar deneyiniz."
                return render_template("index.html", error_headline=error_headline, error_text=error_text)

            filename = secure_filename(file.filename)
            
            s3_client.upload_fileobj(file, "ceveri", filename)

            source_language_abb = request.form.get('source-language')

            if (source_language_abb == "ml") or (source_language_abb == "un"):
                detect_source_language = True
            
            else:
                detect_source_language = False

            extension = filename.rsplit(".", 1)[1]
            
            with open(filename, 'wb') as f:
                s3_client.download_fileobj('ceveri', filename, f)

            be_translation = time.time()
            tr_file_path, chars_translated, source_languages = translate(extension, filename, source_language_abb, detect_source_language)
            af_translation = time.time()
            tr_time = af_translation - be_translation
            
            saving_usd = round((chars_translated / 1000000) * 20 * 2, 3)
        
            if detect_source_language:
                saving_usd *= 2    
                for source_language in source_languages:
                    try:
                        source_languages[source_languages.index(source_language)] = language_abbs[source_language]
                    except:
                        pass

            else:
                source_languages = [language_abbs[source_language_abb]]

            tr_file_path = tr_file_path.split("\\")[-1]
            return render_template("index.html", 
                                    download_file_name="download/"+tr_file_path, 
                                    source_languages=source_languages, 
                                    ready_to_download=True,
                                    chars_translated=chars_translated,
                                    tr_time=tr_time,
                                    saving_usd=saving_usd,
                                    saving_try=round(saving_usd*18, 3))

    return render_template("index.html", error_headline="", error_text="")

@app.route("/ekip")
def ekip():
    return render_template("ekip.html")

@app.route("/klasikler")
def klasikler():
    return render_template("klasikler.html")

@app.route("/hedefler")
def hedefler():
    return render_template("hedefler.html")

@app.route("/rehber")
def rehber():
    return render_template("rehber.html")

@app.route("/download/<download_file_name>")
def download(download_file_name):
    file = s3_client.get_object(Bucket='ceveri', Key=download_file_name)
    return Response(
        file['Body'].read(),
        mimetype='text/plain',
        headers={"Content-Disposition": "attachment;filename="+download_file_name}
    )

if __name__ == "__main__":
    app.run(debug=True)